"""Extract raw text from many source types: pdf, docx, txt, csv, json, url.

Parsing works directly from in-memory bytes (no local files are written).

PDFs are parsed page-by-page. Each page first tries embedded text; if that is
too short (scanned / image-only page) it falls back to OCR via pytesseract.
The richer ``parse_pdf_document`` returns a ``ParsedDocument`` so callers can
record page-level provenance, OCR usage and warnings. ``parse_bytes`` stays a
plain ``str`` -> str helper for the non-PDF parsers and for backwards compat.
"""
from __future__ import annotations
import io
import json
import logging
import csv as _csv
from dataclasses import dataclass, field

from config import settings

logger = logging.getLogger("ielts.ingestion")


# --- Rich PDF result types --------------------------------------------------
@dataclass
class ParsedPage:
    pageNumber: int
    text: str
    method: str  # "pdf_text" | "ocr" | "empty"
    charCount: int
    warning: str | None = None


@dataclass
class ParsedDocument:
    text: str
    pages: list[ParsedPage] = field(default_factory=list)
    pageCount: int = 0
    extractedPages: int = 0  # pages that yielded text via embedded pdf text
    ocrPages: int = 0
    emptyPages: int = 0
    extractionMethod: str = "pdf_text"  # "pdf_text" | "ocr" | "mixed" | "empty"
    warnings: list[str] = field(default_factory=list)


class OcrUnavailable(RuntimeError):
    """Raised when OCR is needed but Tesseract / pytesseract is not available."""


def parse_bytes(data: bytes, kind: str | None = None) -> str:
    """Extract text from raw file bytes. Preferred entry point (no disk I/O).

    For PDFs this returns the concatenated text only (embedded + OCR fallback);
    use :func:`parse_pdf_document` when page-level detail is needed.
    """
    kind = (kind or "txt").lower().lstrip(".")
    if kind == "pdf":
        return parse_pdf_document(data).text
    if kind in ("docx", "doc"):
        return _parse_docx_bytes(data)
    if kind == "csv":
        return _parse_csv_bytes(data)
    if kind == "json":
        return _parse_json_bytes(data)
    return data.decode("utf-8", errors="ignore")


# --- OCR helpers ------------------------------------------------------------
def ocr_available() -> bool:
    """True when pytesseract + Pillow import and the tesseract binary is found."""
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except Exception:
        return False
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _ocr_pixmap(pix, language: str) -> str:
    import pytesseract
    from PIL import Image

    img = Image.open(io.BytesIO(pix.tobytes("png")))
    return pytesseract.image_to_string(img, lang=language)


def parse_pdf_document(
    data: bytes,
    *,
    ocr_enabled: bool | None = None,
    dpi: int | None = None,
    min_chars: int | None = None,
    language: str | None = None,
    max_pages: int | None = None,
) -> ParsedDocument:
    """Parse a PDF into a :class:`ParsedDocument`.

    Per page: try ``page.get_text("text")`` first. If the embedded text is
    shorter than ``min_chars`` the page is treated as scanned/image and run
    through OCR (when enabled and available). Each page records its method,
    char count and any warning.
    """
    import fitz  # PyMuPDF

    ocr_enabled = settings.ocr_enabled if ocr_enabled is None else ocr_enabled
    dpi = settings.ocr_dpi if dpi is None else dpi
    min_chars = settings.ocr_min_chars_per_page if min_chars is None else min_chars
    language = settings.ocr_language if language is None else language
    if max_pages is None:
        max_pages = settings.ocr_max_pages  # 0 / None => all pages

    scale = max(0.5, dpi / 72.0)
    pages: list[ParsedPage] = []
    warnings: list[str] = []
    ocr_ready: bool | None = None  # lazily probed only when first needed

    with fitz.open(stream=data, filetype="pdf") as doc:
        page_count = doc.page_count
        limit = page_count
        if max_pages and max_pages > 0:
            limit = min(page_count, max_pages)
            if limit < page_count:
                warnings.append(
                    f"OCR_MAX_PAGES={max_pages}: only the first {limit} of "
                    f"{page_count} pages were processed."
                )
        for i in range(limit):
            page = doc.load_page(i)
            embedded = page.get_text("text") or ""
            if len(embedded.strip()) >= min_chars:
                pages.append(ParsedPage(i + 1, embedded, "pdf_text", len(embedded)))
                continue

            # Page looks scanned / image-only -> OCR fallback.
            if not ocr_enabled:
                pages.append(ParsedPage(
                    i + 1, embedded, "empty", len(embedded),
                    warning="Low embedded text and OCR is disabled (OCR_ENABLED=false).",
                ))
                continue
            if ocr_ready is None:
                ocr_ready = ocr_available()
                if not ocr_ready:
                    warnings.append(
                        "Tesseract is not installed or OCR is disabled. "
                        "Scanned pages could not be read."
                    )
            if not ocr_ready:
                pages.append(ParsedPage(
                    i + 1, embedded, "empty", len(embedded),
                    warning="Tesseract is not installed or OCR is disabled.",
                ))
                continue
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
                ocr_text = _ocr_pixmap(pix, language)
            except Exception as exc:  # noqa: BLE001
                logger.warning("OCR failed on page %s: %s", i + 1, exc)
                pages.append(ParsedPage(
                    i + 1, embedded, "empty", len(embedded),
                    warning=f"OCR failed: {exc}",
                ))
                continue
            if len(ocr_text.strip()) >= min_chars or len(ocr_text.strip()) > len(embedded.strip()):
                pages.append(ParsedPage(i + 1, ocr_text, "ocr", len(ocr_text)))
            else:
                # combine whatever we got; mark empty if essentially nothing
                combined = (embedded + "\n" + ocr_text).strip()
                method = "ocr" if ocr_text.strip() else "empty"
                warning = None if ocr_text.strip() else "No text could be extracted from this page."
                pages.append(ParsedPage(i + 1, combined, method, len(combined), warning=warning))

    extracted = sum(1 for p in pages if p.method == "pdf_text")
    ocr_pages = sum(1 for p in pages if p.method == "ocr")
    empty_pages = sum(1 for p in pages if p.method == "empty")
    if ocr_pages and extracted:
        method = "mixed"
    elif ocr_pages:
        method = "ocr"
    elif extracted:
        method = "pdf_text"
    else:
        method = "empty"

    full_text = "\n".join(p.text for p in pages if p.text.strip())
    return ParsedDocument(
        text=full_text, pages=pages, pageCount=page_count,
        extractedPages=extracted, ocrPages=ocr_pages, emptyPages=empty_pages,
        extractionMethod=method, warnings=warnings,
    )


def _parse_pdf_bytes(data: bytes) -> str:
    """Backwards-compatible thin wrapper returning concatenated PDF text."""
    return parse_pdf_document(data).text


def _parse_docx_bytes(data: bytes) -> str:
    import docx

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _parse_csv_bytes(data: bytes) -> str:
    text = data.decode("utf-8", errors="ignore")
    try:
        import pandas as pd

        df = pd.read_csv(io.StringIO(text), on_bad_lines="skip")
        return df.to_string(index=False)
    except ImportError:
        rows = []
        for row in _csv.reader(io.StringIO(text)):
            rows.append(" ".join(cell for cell in row if cell))
        return "\n".join(rows)
    except Exception as exc:
        logger.warning("CSV parse failed (%s); falling back to raw text", exc)
        return text


def _parse_json_bytes(data: bytes) -> str:
    try:
        obj = json.loads(data.decode("utf-8", errors="ignore"))
    except Exception:
        return data.decode("utf-8", errors="ignore")
    return _flatten_json(obj)


def _flatten_json(data) -> str:
    out: list[str] = []

    def walk(node):
        if isinstance(node, dict):
            for v in node.values():
                walk(v)
        elif isinstance(node, list):
            for v in node:
                walk(v)
        elif isinstance(node, str):
            out.append(node)
        elif node is not None:
            out.append(str(node))

    walk(data)
    return "\n".join(out)


async def parse_url(url: str) -> tuple[str, str]:
    """Fetch a URL and return (title, text)."""
    import httpx
    from bs4 import BeautifulSoup

    headers = {"User-Agent": "Mozilla/5.0 (IELTS-AI-Mastery; +local)"}
    async with httpx.AsyncClient(timeout=30, follow_redirects=True, headers=headers) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        html = resp.text
    # Use lxml if installed (faster), else the stdlib html.parser.
    try:
        soup = BeautifulSoup(html, "lxml")
    except Exception:
        soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()
    title = (soup.title.string.strip() if soup.title and soup.title.string else url)
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    return title, text
